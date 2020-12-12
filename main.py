import cv2
import pytesseract
import numpy as np
import docx
from docx import Document 

pytesseract.pytesseract.tesseract_cmd = r'D:\Tesseract-OCR\tesseract'

img_cv = cv2.imread(r'C:\Users\chama\Desktop\Capture.png')

img_rgb = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
# cv2.imshow("i", img_rgb)
# cv2.waitKey()
t = pytesseract.image_to_string(img_rgb, lang='sin')

pdf = pytesseract.image_to_pdf_or_hocr(img_rgb, lang='sin', extension='pdf')
with open('test.pdf', 'w+b') as f:
    f.write(pdf)

print(t)
# text_file = open("Output.txt", "w")
# print(t, file=text_file)

document = Document() 

# styles_element = document.styles.element
# rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
# lang_default = rpr_default.xpath('w:lang')[0]
# lang_default.set(docx.oxml.shared.qn('w:val'),'de-DE')

# document.add_paragraph( 
#     t
# ) 
# document.save('demo.docx') 

import re

def sanitize_str(s):
    control_chars = "\x00-\x1f\x7f-\x9f"
    control_char_re = re.compile("[%s]" % control_chars)
    return control_char_re.sub("", s)

from docx.shared import Pt

style = document.styles['Normal']
font = style.font
font.name = 'Iskoola Pota'
font.size = Pt(10)

document.add_paragraph(sanitize_str(t))

document.save('demo.docx') 