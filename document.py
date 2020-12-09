from docx import Document 

document = Document() 
document.add_heading('Document Title', 0) 
document.add_heading('Heading, level 1', level=1) 
document.add_paragraph('Intense quote', style='Intense Quote') 
 
document.add_paragraph( 
    'first item in unordered list', style='List Bullet' 
) 

document.save('demo.docx') 